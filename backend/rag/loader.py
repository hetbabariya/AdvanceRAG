from __future__ import annotations

from dataclasses import dataclass
from datetime import datetime
import hashlib
import os
import re
import unicodedata
from pathlib import Path
from typing import Dict, List, Optional, Tuple

import pymupdf
from bs4 import BeautifulSoup
from docx import Document as DocxDocument
from langchain_core.documents import Document
from langchain_community.document_loaders import WebBaseLoader
from langchain_text_splitters import RecursiveCharacterTextSplitter


try:
    from langsmith import traceable  # type: ignore
except Exception:  # pragma: no cover
    def traceable(*_args, **_kwargs):  # type: ignore
        def _decorator(fn):
            return fn

        return _decorator


class OptimizedPreprocessedLoader:
    def __init__(
        self,
        chunk_size: int = 500,
        chunk_overlap: int = 100,
        preprocessing_config: Optional[Dict[str, bool]] = None,
    ):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap

        self.splitter = RecursiveCharacterTextSplitter.from_tiktoken_encoder(
            model_name="gpt-4",
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
        )

        self.config: Dict[str, bool] = {
            "remove_extra_whitespace": True,
            "remove_extra_newlines": True,
            "normalize_unicode": True,
            "fix_encoding_errors": True,
            "fix_common_ocr_errors": False,
            "remove_headers_footers": False,
            "remove_page_numbers": False,
            "remove_short_lines": False,
        }

        if preprocessing_config:
            self.config.update(preprocessing_config)

    def load_and_split(self, source: str, custom_metadata: Optional[Dict] = None) -> List[Document]:
        if source.startswith(("http://", "https://")):
            return self._load_from_url(source, custom_metadata or {})
        return self._load_from_file(source, custom_metadata or {})

    def load_and_split_youtube(
        self,
        url: str,
        custom_metadata: Optional[Dict] = None,
        languages: Optional[List[str]] = None,
    ) -> Tuple[str, List[Document]]:
        from urllib.parse import parse_qs, urlparse

        from xml.etree.ElementTree import ParseError

        from youtube_transcript_api import YouTubeTranscriptApi

        custom_metadata = custom_metadata or {}
        languages = languages or ["en", "hi", "en-US", "hi-IN"]

        parsed = urlparse(url)
        host = (parsed.netloc or "").lower()
        if host.startswith("www."):
            host = host[4:]

        video_id = ""
        if host.endswith("youtu.be"):
            video_id = parsed.path.strip("/")
        else:
            qs = parse_qs(parsed.query)
            video_id = (qs.get("v") or [""])[0]

        video_id = (video_id or "").strip()
        if not video_id:
            raise ValueError("Could not extract YouTube video id")

        transcript_items = None
        last_error: Exception | None = None

        for _ in range(2):
            try:
                yt = YouTubeTranscriptApi()
                transcript_items = yt.fetch(video_id, languages=languages)
                break
            except TypeError:
                transcript_items = YouTubeTranscriptApi.get_transcript(video_id)
                break
            except (ParseError, Exception) as e:
                last_error = e

        if not transcript_items:
            if last_error is not None:
                raise ValueError(f"Failed to fetch transcript: {last_error}")
            raise ValueError("No transcript available for this video")

        lines: List[str] = []
        for snippet in transcript_items:
            if isinstance(snippet, dict):
                text = str(snippet.get("text") or "").strip()
            else:
                text = str(getattr(snippet, "text", "") or "").strip()
            if text:
                lines.append(text)

        transcript_text = "\n".join(lines)
        transcript_text = self.preprocess_text(transcript_text)
        if not transcript_text.strip():
            raise ValueError("No transcript available for this video")

        display_name = f"youtube_{video_id}.txt"
        doc_id = self._generate_doc_id(url)

        base_doc = Document(
            page_content=transcript_text,
            metadata={
                "source": url,
                "source_type": "youtube",
                "file_type": "youtube",
                "file_name": display_name,
                "video_id": video_id,
                "loaded_at": datetime.now().isoformat(),
                "doc_id": doc_id,
                **custom_metadata,
            },
        )

        chunks = self.splitter.split_documents([base_doc])
        for i, chunk in enumerate(chunks):
            chunk.metadata["chunk_index"] = i
            chunk.metadata["total_chunks"] = len(chunks)
            chunk.metadata["chunk_id"] = f"{doc_id}_chunk_{i}"
            chunk.metadata["chunk_size"] = len(chunk.page_content)

        return display_name, chunks

    load_and_split_youtube = traceable(name="loader.youtube")(load_and_split_youtube)

    def load_and_split_file(self, file_path: str, original_name: str, custom_metadata: Optional[Dict] = None) -> List[Document]:
        docs = self._load_from_file(file_path, custom_metadata or {})
        for d in docs:
            d.metadata["file_name"] = original_name
        return docs

    def _load_from_url(self, url: str, custom_metadata: Dict) -> List[Document]:
        loader = WebBaseLoader(web_paths=[url])
        docs = loader.load()

        for doc in docs:
            doc.page_content = self.preprocess_text(doc.page_content)
            doc.metadata.update(
                {
                    "source": url,
                    "source_type": "url",
                    "file_type": "html",
                    "loaded_at": datetime.now().isoformat(),
                    "doc_id": self._generate_doc_id(url),
                    **custom_metadata,
                }
            )

        return self._optimized_split(docs)

    def _chunking_strategy(self) -> str:
        raw = (os.getenv("CHUNK_STRATEGY", "") or "").strip().lower()
        return raw if raw in {"structural", "flat"} else "structural"

    def _load_from_file(self, file_path: str, custom_metadata: Dict) -> List[Document]:
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        suffix = path.suffix.lower()

        structural = False
        if suffix == ".pdf":
            page_contents, file_metadata = self._extract_pdf_structural(path)
            structural = True
        elif suffix == ".docx":
            page_contents, file_metadata = self._extract_docx_with_sections(path)
        elif suffix == ".txt":
            page_contents, file_metadata = self._extract_txt(path)
        elif suffix == ".md":
            page_contents, file_metadata = self._extract_md(path)
            structural = True
        else:
            raise ValueError(f"Unsupported format: {suffix}")

        base_metadata = {
            "source": str(path.absolute()),
            "source_type": "file",
            "file_type": suffix[1:],
            "file_name": path.name,
            "file_size": path.stat().st_size,
            "loaded_at": datetime.now().isoformat(),
            "doc_id": self._generate_doc_id(str(path)),
            **file_metadata,
            **custom_metadata,
        }

        if structural and self._chunking_strategy() == "structural":
            cleaned_pages: List[Dict] = []
            for page_info in page_contents:
                text = self.preprocess_structural_text(page_info["text"])
                if not text.strip():
                    continue
                cleaned_pages.append({"page_num": page_info["page_num"], "text": text})
            return self._structural_chunks(cleaned_pages, base_metadata)

        docs_with_pages: List[Document] = []
        for page_info in page_contents:
            preprocessed_text = self.preprocess_text(page_info["text"])
            if not preprocessed_text.strip():
                continue

            docs_with_pages.append(
                Document(
                    page_content=preprocessed_text,
                    metadata={
                        **base_metadata,
                        "page_number": page_info["page_num"],
                        "page_start": page_info["page_num"],
                        "page_end": page_info["page_num"],
                    },
                )
            )

        chunks = self._optimized_split_with_pages(docs_with_pages)
        return chunks

    # ------------------------------------------------------------------
    # Structure-aware parsing & chunking (PDF / Markdown)
    # ------------------------------------------------------------------

    _HEADING_RE = re.compile(r"^(#{1,6})\s+(.*)\s*$")
    MIN_SECTION_TOKENS = 40

    def preprocess_structural_text(self, text: str) -> str:
        """Light-touch cleanup for structure-aware parsing.

        Applies unicode/encoding fixes only — whitespace collapsing would
        destroy the markdown layout that heading detection depends on.
        """
        if not text:
            return ""
        if self.config.get("fix_encoding_errors", False):
            text = self._fix_encoding_errors(text)
        if self.config.get("normalize_unicode", False):
            text = unicodedata.normalize("NFKC", text)
        return text

    def _extract_pdf_structural(self, path: Path) -> Tuple[List[Dict], Dict]:
        """Extract PDF as per-page markdown via pymupdf4llm (layout-aware).

        Falls back to raw page.get_text() if pymupdf4llm is unavailable or fails.
        """
        try:
            import pymupdf4llm

            doc = pymupdf.open(path)
            md_pages = pymupdf4llm.to_markdown(doc, page_chunks=True)

            metadata: Dict = {
                "parser": "pymupdf4llm",
                "page_count": len(doc),
                "author": doc.metadata.get("author", ""),
                "title": doc.metadata.get("title", ""),
                "subject": doc.metadata.get("subject", ""),
                "keywords": doc.metadata.get("keywords", ""),
                "creator": doc.metadata.get("creator", ""),
                "creation_date": doc.metadata.get("creationDate", ""),
                "modification_date": doc.metadata.get("modDate", ""),
            }
            doc.close()

            page_contents: List[Dict] = []
            for idx, item in enumerate(md_pages or []):
                item = item if isinstance(item, dict) else {}
                text = str(item.get("text") or "")
                meta = item.get("metadata") or {}
                try:
                    page_num = int(meta.get("page", idx)) + 1  # 0-based → 1-based
                except (TypeError, ValueError):
                    page_num = idx + 1
                if text.strip():
                    page_contents.append({"page_num": max(1, page_num), "text": text})

            return page_contents, {k: v for k, v in metadata.items() if v}
        except Exception:
            logger.exception(
                "pymupdf4llm extraction failed for %s; falling back to raw PyMuPDF", path.name
            )
            return self._extract_pdf_with_pages(path)

    def _split_sections(self, md: str) -> List[Dict[str, object]]:
        """Split markdown into heading-delimited sections with char offsets."""
        sections: List[Dict[str, object]] = []
        lines = md.split("\n")
        total_len = len(md)

        line_offsets: List[int] = [0] * len(lines)
        acc = 0
        for i, ln in enumerate(lines):
            line_offsets[i] = acc
            acc += len(ln) + 1

        def _flush(start_line: int, end_line: int, title: str, level: int) -> None:
            if start_line >= end_line:
                return
            content = "\n".join(lines[start_line:end_line]).strip()
            if not content:
                return
            sections.append({
                "title": title,
                "level": level,
                "content": content,
                "start_offset": line_offsets[start_line],
                "end_offset": line_offsets[end_line] if end_line < len(lines) else total_len,
            })

        cur_title, cur_level, cur_start = "", 0, 0
        in_code = False
        for i, line in enumerate(lines):
            if line.strip().startswith("```"):
                in_code = not in_code
                continue
            if in_code:
                continue
            m = self._HEADING_RE.match(line)
            if m:
                _flush(cur_start, i, cur_title, cur_level)
                cur_title = (m.group(2) or "").strip()
                cur_level = len(m.group(1))
                cur_start = i
        _flush(cur_start, len(lines), cur_title, cur_level)

        if not sections and md.strip():
            return [{
                "title": "",
                "level": 0,
                "content": md.strip(),
                "start_offset": 0,
                "end_offset": total_len,
            }]
        return sections

    def _merge_tiny_sections(self, sections: List[Dict[str, object]]) -> List[Dict[str, object]]:
        """Absorb sections smaller than MIN_SECTION_TOKENS into neighbours so
        we don't emit one-line chunks for short headings."""
        if len(sections) <= 1:
            return sections

        merged: List[Dict[str, object]] = []
        for sec in sections:
            if merged and self._token_len(str(sec["content"])) < self.MIN_SECTION_TOKENS:
                prev = merged[-1]
                prev["content"] = f"{prev['content']}\n\n{sec['content']}".strip()
                prev["end_offset"] = sec["end_offset"]
            elif (
                merged
                and self._token_len(str(merged[-1]["content"])) < self.MIN_SECTION_TOKENS
            ):
                prev = merged[-1]
                prev["content"] = f"{prev['content']}\n\n{sec['content']}".strip()
                prev["end_offset"] = sec["end_offset"]
            else:
                merged.append(dict(sec))
        return merged

    def _structural_chunks(self, pages: List[Dict], base_metadata: Dict) -> List[Document]:
        """Section-aware chunking across page boundaries.

        Every chunk keeps its section title (prepended to the text) so the
        embedding carries topical context even when a section is split.
        """
        parts: List[str] = []
        page_marks: List[Tuple[int, int]] = []
        cursor = 0
        for p in pages:
            text = p["text"].rstrip()
            if not text.strip():
                continue
            page_marks.append((cursor, int(p["page_num"])))
            parts.append(text)
            cursor += len(text) + 2  # "\n\n" joiner
        full_md = "\n\n".join(parts)

        if not full_md.strip():
            return []

        def _page_at(offset: int) -> int:
            pg = page_marks[0][1] if page_marks else 1
            for start, num in page_marks:
                if offset >= start:
                    pg = num
                else:
                    break
            return pg

        sections = self._merge_tiny_sections(self._split_sections(full_md))

        final_chunks: List[Document] = []
        chunk_counter = 0
        doc_id = str(base_metadata.get("doc_id", "unknown"))

        for si, sec in enumerate(sections):
            content = str(sec["content"]).strip()
            if not content:
                continue
            title = str(sec["title"] or "")
            prefix = f"{title}\n\n" if title else ""

            start_off = int(sec["start_offset"])
            end_off = int(sec["end_offset"])
            p_start = _page_at(start_off)
            p_end = _page_at(max(end_off - 1, start_off))
            page_label = str(p_start) if p_start == p_end else f"{p_start}-{p_end}"

            if self._token_len(prefix + content) <= self.chunk_size:
                pieces = [prefix + content]
            else:
                pieces = [f"{prefix}{piece}" for piece in self.splitter.split_text(content)]

            for piece in pieces:
                final_chunks.append(
                    Document(
                        page_content=piece,
                        metadata={
                            **base_metadata,
                            "section_index": si,
                            "section_title": title,
                            "section_level": int(sec["level"]),
                            "section_count": len(sections),
                            "page_number": page_label,
                            "page_start": p_start,
                            "page_end": p_end,
                            "chunk_index": chunk_counter,
                            "chunk_id": f"{doc_id}_chunk_{chunk_counter}",
                            "chunk_size": len(piece),
                        },
                    )
                )
                chunk_counter += 1

        for chunk in final_chunks:
            chunk.metadata["total_chunks"] = len(final_chunks)
        return final_chunks

    def _token_len(self, text: str) -> int:
        try:
            import tiktoken

            enc = tiktoken.get_encoding("cl100k_base")
            return len(enc.encode(text or ""))
        except Exception:
            return len((text or "").split())

    def _extract_pdf_with_pages(self, path: Path) -> Tuple[List[Dict], Dict]:
        doc = pymupdf.open(path)
        page_contents: List[Dict] = []

        for page_num, page in enumerate(doc, start=1):
            page_text = page.get_text()
            if page_text.strip():
                page_contents.append({"page_num": page_num, "text": page_text})

        metadata = {
            "page_count": len(doc),
            "author": doc.metadata.get("author", ""),
            "title": doc.metadata.get("title", ""),
            "subject": doc.metadata.get("subject", ""),
            "keywords": doc.metadata.get("keywords", ""),
            "creator": doc.metadata.get("creator", ""),
            "creation_date": doc.metadata.get("creationDate", ""),
            "modification_date": doc.metadata.get("modDate", ""),
        }

        doc.close()
        metadata = {k: v for k, v in metadata.items() if v}
        return page_contents, metadata

    def _extract_md(self, path: Path) -> Tuple[List[Dict], Dict]:
        with open(path, "r", encoding="utf-8", errors="ignore") as f:
            text_content = f.read()

        metadata = {
            "character_count": len(text_content),
        }
        return [{"page_num": 1, "text": text_content}], metadata

    def _extract_docx_with_sections(self, path: Path) -> Tuple[List[Dict], Dict]:
        doc = DocxDocument(path)
        words_per_page = 500

        page_contents: List[Dict] = []
        current_page_text: List[str] = []
        current_word_count = 0
        page_num = 1

        for para in doc.paragraphs:
            if not para.text.strip():
                continue
            para_words = len(para.text.split())
            current_page_text.append(para.text)
            current_word_count += para_words

            if current_word_count >= words_per_page:
                page_contents.append({"page_num": page_num, "text": "\n\n".join(current_page_text)})
                current_page_text = []
                current_word_count = 0
                page_num += 1

        if current_page_text:
            page_contents.append({"page_num": page_num, "text": "\n\n".join(current_page_text)})

        for table in doc.tables:
            page_num += 1
            table_text: List[str] = []
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells)
                if row_text.strip():
                    table_text.append(row_text)

            if table_text:
                page_contents.append({"page_num": page_num, "text": "\n".join(table_text)})

        core_props = doc.core_properties
        metadata = {
            "author": core_props.author or "",
            "title": core_props.title or "",
            "subject": core_props.subject or "",
            "keywords": core_props.keywords or "",
            "created": core_props.created.isoformat() if core_props.created else "",
            "modified": core_props.modified.isoformat() if core_props.modified else "",
            "paragraph_count": len(doc.paragraphs),
            "table_count": len(doc.tables),
            "estimated_pages": len(page_contents),
        }

        metadata = {k: v for k, v in metadata.items() if v}
        return page_contents, metadata

    def _extract_txt(self, path: Path) -> Tuple[List[Dict], Dict]:
        with open(path, "r", encoding="utf-8", errors="ignore") as f:
            text_content = f.read()

        words_per_page = 500
        words = text_content.split()

        page_contents: List[Dict] = []
        page_num = 1
        current_word_count = 0
        current_page_words: List[str] = []

        for word in words:
            current_page_words.append(word)
            current_word_count += 1

            if current_word_count >= words_per_page:
                page_contents.append({
                    "page_num": page_num,
                    "text": " ".join(current_page_words)
                })
                current_page_words = []
                current_word_count = 0
                page_num += 1

        if current_page_words:
            page_contents.append({
                "page_num": page_num,
                "text": " ".join(current_page_words)
            })

        if not page_contents:
            page_contents.append({"page_num": 1, "text": text_content})

        metadata = {
            "estimated_pages": len(page_contents),
            "word_count": len(words),
            "character_count": len(text_content),
        }

        return page_contents, metadata

    def _extract_html_with_metadata(self, path: Path) -> Tuple[str, Dict]:
        with open(path, "r", encoding="utf-8") as f:
            html_content = f.read()

        soup = BeautifulSoup(html_content, "html.parser")

        metadata = {
            "title": soup.title.string if soup.title else "",
            "description": "",
            "keywords": "",
            "author": "",
        }

        for meta in soup.find_all("meta"):
            name = (meta.get("name", "") or "").lower()
            content = meta.get("content", "")

            if name == "description":
                metadata["description"] = content
            elif name == "keywords":
                metadata["keywords"] = content
            elif name == "author":
                metadata["author"] = content

        for tag in soup(["script", "style", "nav", "footer", "header"]):
            tag.decompose()

        text = soup.get_text(separator="\n", strip=True)
        metadata = {k: v for k, v in metadata.items() if v}
        return text, metadata

    def _optimized_split(self, docs: List[Document]) -> List[Document]:
        final_chunks: List[Document] = []
        chunk_counter = 0

        for doc in docs:
            section_metadata = dict(doc.metadata)
            if self._token_len(doc.page_content) <= self.chunk_size:
                chunks = [doc.page_content]
            else:
                chunks = self.splitter.split_text(doc.page_content)

            for chunk in chunks:
                final_chunks.append(
                    Document(
                        page_content=chunk,
                        metadata={
                            **section_metadata,
                            "chunk_index": chunk_counter,
                            "chunk_id": f"{section_metadata.get('doc_id', 'unknown')}_chunk_{chunk_counter}",
                            "chunk_size": len(chunk),
                        },
                    )
                )
                chunk_counter += 1

        for chunk in final_chunks:
            chunk.metadata["total_chunks"] = len(final_chunks)

        return final_chunks

    def _optimized_split_with_pages(self, docs_with_pages: List[Document]) -> List[Document]:
        final_chunks: List[Document] = []
        chunk_counter = 0

        for doc in docs_with_pages:
            page_metadata = dict(doc.metadata)
            if self._token_len(doc.page_content) <= self.chunk_size:
                chunks = [doc.page_content]
            else:
                chunks = self.splitter.split_text(doc.page_content)

            for chunk in chunks:
                final_chunks.append(
                    Document(
                        page_content=chunk,
                        metadata={
                            **page_metadata,
                            "chunk_index": chunk_counter,
                            "chunk_id": f"{page_metadata.get('doc_id', 'unknown')}_chunk_{chunk_counter}",
                            "chunk_size": len(chunk),
                        },
                    )
                )
                chunk_counter += 1

        for chunk in final_chunks:
            chunk.metadata["total_chunks"] = len(final_chunks)

        return final_chunks

    def preprocess_text(self, text: str) -> str:
        if not text:
            return ""

        if self.config.get("fix_encoding_errors", False):
            text = self._fix_encoding_errors(text)

        if self.config.get("normalize_unicode", False):
            text = unicodedata.normalize("NFKC", text)

        if self.config.get("fix_common_ocr_errors", False):
            text = self._fix_ocr_errors(text)

        if self.config.get("remove_headers_footers", False):
            text = self._remove_headers_footers(text)

        if self.config.get("remove_page_numbers", False):
            text = self._remove_page_numbers(text)

        if self.config.get("remove_extra_whitespace", False):
            text = re.sub(r"[ \t]+", " ", text)

        if self.config.get("remove_extra_newlines", False):
            text = re.sub(r"\n{3,}", "\n\n", text)

        if self.config.get("remove_short_lines", False):
            lines = text.split("\n")
            lines = [line for line in lines if len(line.split()) >= 3 or line.strip() == ""]
            text = "\n".join(lines)

        return text.strip()

    def _fix_encoding_errors(self, text: str) -> str:
        replacements = {
            "â€™": "'",
            "â€œ": '"',
            "â€": '"',
            "â€\"": "—",
            "Â": "",
            "\x00": "",
            "\ufffd": "",
        }
        for old, new in replacements.items():
            text = text.replace(old, new)
        return text

    def _fix_ocr_errors(self, text: str) -> str:
        return re.sub(r"(\w+)-\n(\w+)", r"\1\2", text)

    def _remove_headers_footers(self, text: str) -> str:
        lines = text.split("\n")
        cleaned: List[str] = []

        skip_patterns = [
            r"^page \d+",
            r"^\d+ of \d+$",
            r"^chapter \d+",
            r"^confidential",
            r"^\d+$",
        ]

        for line in lines:
            line_lower = line.lower().strip()
            if any(re.match(p, line_lower) for p in skip_patterns):
                continue
            cleaned.append(line)

        return "\n".join(cleaned)

    def _remove_page_numbers(self, text: str) -> str:
        text = re.sub(r"^\s*\d+\s*$", "", text, flags=re.MULTILINE)
        text = re.sub(r"\bPage\s+\d+\b", "", text, flags=re.IGNORECASE)
        return text

    def _generate_doc_id(self, source: str) -> str:
        return hashlib.md5(source.encode()).hexdigest()[:16]
